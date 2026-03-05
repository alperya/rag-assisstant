import os
import uuid
from typing import Optional

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
import pypdf
import docx

from app.config import get_settings


class DocumentProcessor:
    """Handles parsing and chunking of PDF, DOCX, and TXT documents."""

    def __init__(self):
        settings = get_settings()
        self.default_chunk_size = settings.chunk_size
        self.default_chunk_overlap = settings.chunk_overlap

    def _get_splitter(self, chunk_size: int | None = None, chunk_overlap: int | None = None):
        """Create a text splitter with given or default params."""
        cs = chunk_size if chunk_size and chunk_size > 0 else self.default_chunk_size
        co = chunk_overlap if chunk_overlap and chunk_overlap >= 0 else self.default_chunk_overlap
        return RecursiveCharacterTextSplitter(
            chunk_size=cs,
            chunk_overlap=co,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""],
        )

    def process_pdf(self, file_path: str, filename: str) -> tuple[list[Document], int]:
        """Extract text from PDF with page-level tracking."""
        documents = []
        reader = pypdf.PdfReader(file_path)
        page_count = len(reader.pages)

        for page_num, page in enumerate(reader.pages, start=1):
            text = page.extract_text()
            if text and text.strip():
                documents.append(
                    Document(
                        page_content=text.strip(),
                        metadata={
                            "source": filename,
                            "page": page_num,
                            "total_pages": page_count,
                            "file_type": "pdf",
                        },
                    )
                )

        return documents, page_count

    def process_docx(self, file_path: str, filename: str) -> tuple[list[Document], Optional[int]]:
        """Extract text from DOCX files."""
        doc = docx.Document(file_path)
        documents = []
        current_text = []
        page_estimate = 1

        for para in doc.paragraphs:
            if para.text.strip():
                current_text.append(para.text.strip())

            # Estimate page breaks (~3000 chars per page)
            full_text = "\n".join(current_text)
            if len(full_text) > 3000:
                documents.append(
                    Document(
                        page_content=full_text,
                        metadata={
                            "source": filename,
                            "page": page_estimate,
                            "file_type": "docx",
                        },
                    )
                )
                current_text = []
                page_estimate += 1

        # Remaining text
        if current_text:
            documents.append(
                Document(
                    page_content="\n".join(current_text),
                    metadata={
                        "source": filename,
                        "page": page_estimate,
                        "file_type": "docx",
                    },
                )
            )

        return documents, page_estimate

    def process_txt(self, file_path: str, filename: str) -> tuple[list[Document], Optional[int]]:
        """Extract text from TXT files."""
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()

        # Split into estimated pages (~3000 chars each)
        page_size = 3000
        pages = [text[i : i + page_size] for i in range(0, len(text), page_size)]
        documents = []

        for page_num, page_text in enumerate(pages, start=1):
            if page_text.strip():
                documents.append(
                    Document(
                        page_content=page_text.strip(),
                        metadata={
                            "source": filename,
                            "page": page_num,
                            "file_type": "txt",
                        },
                    )
                )

        return documents, len(pages)

    def process_text_input(self, text: str, title: str) -> tuple[list[Document], int]:
        """Process plain text input from user."""
        page_size = 3000
        pages = [text[i : i + page_size] for i in range(0, len(text), page_size)]
        documents = []

        for page_num, page_text in enumerate(pages, start=1):
            if page_text.strip():
                documents.append(
                    Document(
                        page_content=page_text.strip(),
                        metadata={
                            "source": title,
                            "page": page_num,
                            "file_type": "text",
                        },
                    )
                )

        return documents, len(pages)

    def process_file(
        self,
        file_path: str,
        filename: str,
        chunk_size: int | None = None,
        chunk_overlap: int | None = None,
    ) -> tuple[list[Document], Optional[int], str]:
        """Process a file based on its extension."""
        ext = os.path.splitext(filename)[1].lower()

        if ext == ".pdf":
            raw_docs, page_count = self.process_pdf(file_path, filename)
        elif ext == ".docx":
            raw_docs, page_count = self.process_docx(file_path, filename)
        elif ext == ".txt":
            raw_docs, page_count = self.process_txt(file_path, filename)
        else:
            raise ValueError(f"Unsupported file type: {ext}")

        # Split into chunks while preserving metadata
        splitter = self._get_splitter(chunk_size, chunk_overlap)
        chunks = splitter.split_documents(raw_docs)

        # Ensure each chunk has the document ID
        doc_id = str(uuid.uuid4())
        for chunk in chunks:
            chunk.metadata["doc_id"] = doc_id

        return chunks, page_count, doc_id

    def process_raw_text(
        self,
        text: str,
        title: str,
        chunk_size: int | None = None,
        chunk_overlap: int | None = None,
    ) -> tuple[list[Document], int, str]:
        """Process plain text with chunking."""
        raw_docs, page_count = self.process_text_input(text, title)

        splitter = self._get_splitter(chunk_size, chunk_overlap)
        chunks = splitter.split_documents(raw_docs)

        doc_id = str(uuid.uuid4())
        for chunk in chunks:
            chunk.metadata["doc_id"] = doc_id

        return chunks, page_count, doc_id
